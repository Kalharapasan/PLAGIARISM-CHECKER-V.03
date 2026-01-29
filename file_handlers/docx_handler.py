import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from datetime import datetime
import tempfile
import os

class DOCXHandler:
    NAMESPACES = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'v': 'urn:schemas-microsoft-com:vml'
    }
    
    def __init__(self, config=None):
        self.config = config or {}
        self._register_namespaces()
    
    def _register_namespaces(self):
        for prefix, uri in self.NAMESPACES.items():
            ET.register_namespace(prefix, uri)
    
    def extract_text(self, filepath: str) -> str:
        try:
            return self._extract_with_docx(filepath)
        except ImportError:
            return self._extract_manual(filepath)
        except Exception as e:
            try:
                return self._extract_manual(filepath)
            except Exception as e2:
                raise Exception(f"Failed to extract text from DOCX: {e2}")
    
    def _extract_with_docx(self, filepath: str) -> str:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        
        doc = Document(filepath)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
        footer = section.footer
        for paragraph in footer.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        text_boxes = self._extract_text_boxes_from_docx(filepath)
        text_parts.extend(text_boxes)
        
        return '\n'.join(text_parts)

    def _extract_text_boxes_from_docx(self, filepath: str) -> List[str]:
        text_boxes = []
        try:
            with zipfile.ZipFile(filepath) as docx:
                if 'word/document.xml' not in docx.namelist():
                    return text_boxes
                xml_content = docx.read('word/document.xml').decode('utf-8')
                
                root = ET.fromstring(xml_content)
                for shape in root.findall('.//w:txbxContent', self.NAMESPACES):
                    for paragraph in shape.findall('.//w:p', self.NAMESPACES):
                        text = self._extract_text_from_paragraph(paragraph)
                        if text.strip():
                            text_boxes.append(text)
                
        except Exception as e:
            print(f"Warning: Could not extract text boxes: {e}")
        
        return text_boxes
    
    def _extract_manual(self, filepath: str) -> str:
        text_parts = []
        
        with zipfile.ZipFile(filepath) as docx:
            if 'word/document.xml' in docx.namelist():
                xml_content = docx.read('word/document.xml').decode('utf-8')
                text_parts.append(self._parse_document_xml(xml_content))
            header_texts = self._extract_headers_manual(docx)
            text_parts.extend(header_texts)
            footer_texts = self._extract_footers_manual(docx)
            text_parts.extend(footer_texts)
            footnote_texts = self._extract_footnotes_manual(docx)
            text_parts.extend(footnote_texts)
            endnote_texts = self._extract_endnotes_manual(docx)
            text_parts.extend(endnote_texts)
        return '\n'.join(filter(None, text_parts))

    def _parse_document_xml(self, xml_content: str) -> str:
        try:
            root = ET.fromstring(xml_content)
            paragraphs = []
            for para in root.findall('.//w:p', self.NAMESPACES):
                para_text = self._extract_text_from_paragraph(para)
                if para_text.strip():
                    paragraphs.append(para_text)
            
            return '\n'.join(paragraphs)
        except Exception as e:
            print(f"Warning: Error parsing document XML: {e}")
            return self._extract_text_regex(xml_content)

    def _extract_text_from_paragraph(self, paragraph: ET.Element) -> str:
        text_parts = []
        for run in paragraph.findall('.//w:t', self.NAMESPACES):
            if run.text:
                text_parts.append(run.text)
        for tab in paragraph.findall('.//w:tab', self.NAMESPACES):
            text_parts.append('\t')
        for br in paragraph.findall('.//w:br', self.NAMESPACES):
            text_parts.append('\n')
        
        return ''.join(text_parts)
    
    def _extract_text_regex(self, xml_content: str) -> str:
        text = re.sub(r'<[^>]+>', ' ', xml_content)
        text = re.sub(r'\s+', ' ', text)
        import html
        text = html.unescape(text)
        
        return text.strip()
    
    def _extract_headers_manual(self, docx: zipfile.ZipFile) -> List[str]:
        header_texts = []
        try:
            if 'word/_rels/document.xml.rels' in docx.namelist():
                rels_content = docx.read('word/_rels/document.xml.rels').decode('utf-8')
                rels_root = ET.fromstring(rels_content)
                
                for rel in rels_root.findall('.//Relationship', {'Type': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header'}):
                    header_path = f"word/{rel.attrib['Target']}"
                    
                    if header_path in docx.namelist():
                        header_xml = docx.read(header_path).decode('utf-8')
                        header_text = self._parse_document_xml(header_xml)
                        if header_text.strip():
                            header_texts.append(header_text)
        
        except Exception as e:
            print(f"Warning: Could not extract headers: {e}")
        
        return header_texts
    
    def _extract_footers_manual(self, docx: zipfile.ZipFile) -> List[str]:
        footer_texts = []
        try:
            if 'word/_rels/document.xml.rels' in docx.namelist():
                rels_content = docx.read('word/_rels/document.xml.rels').decode('utf-8')
                rels_root = ET.fromstring(rels_content)
                
                for rel in rels_root.findall('.//Relationship', {'Type': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer'}):
                    footer_path = f"word/{rel.attrib['Target']}"
                    
                    if footer_path in docx.namelist():
                        footer_xml = docx.read(footer_path).decode('utf-8')
                        footer_text = self._parse_document_xml(footer_xml)
                        if footer_text.strip():
                            footer_texts.append(footer_text)
        
        except Exception as e:
            print(f"Warning: Could not extract footers: {e}")
        
        return footer_texts
    
    def _extract_footnotes_manual(self, docx: zipfile.ZipFile) -> List[str]:
        footnote_texts = []
        
        try:
            if 'word/footnotes.xml' in docx.namelist():
                footnotes_xml = docx.read('word/footnotes.xml').decode('utf-8')
                footnote_text = self._parse_document_xml(footnotes_xml)
                if footnote_text.strip():
                    footnote_texts.append(footnote_text)
        
        except Exception as e:
            print(f"Warning: Could not extract footnotes: {e}")
        
        return footnote_texts
    
    def _extract_endnotes_manual(self, docx: zipfile.ZipFile) -> List[str]:
        endnote_texts = []
        
        try:
            if 'word/endnotes.xml' in docx.namelist():
                endnotes_xml = docx.read('word/endnotes.xml').decode('utf-8')
                endnote_text = self._parse_document_xml(endnotes_xml)
                if endnote_text.strip():
                    endnote_texts.append(endnote_text)
        
        except Exception as e:
            print(f"Warning: Could not extract endnotes: {e}")
        
        return endnote_texts
    
    def extract_metadata(self, filepath: str) -> Dict[str, Any]:
        metadata = {
            'filename': Path(filepath).name,
            'file_size': Path(filepath).stat().st_size,
            'modified': datetime.fromtimestamp(Path(filepath).stat().st_mtime).isoformat(),
            'created': datetime.fromtimestamp(Path(filepath).stat().st_ctime).isoformat(),
            'document_properties': {}
        }
        try:
            with zipfile.ZipFile(filepath) as docx:
                if 'docProps/core.xml' in docx.namelist():
                    core_xml = docx.read('docProps/core.xml').decode('utf-8')
                    metadata['document_properties'].update(self._parse_core_properties(core_xml))
                if 'docProps/app.xml' in docx.namelist():
                    app_xml = docx.read('docProps/app.xml').decode('utf-8')
                    metadata['document_properties'].update(self._parse_app_properties(app_xml))
                if 'docProps/custom.xml' in docx.namelist():
                    custom_xml = docx.read('docProps/custom.xml').decode('utf-8')
                    metadata['document_properties'].update(self._parse_custom_properties(custom_xml))
        
        except Exception as e:
            print(f"Warning: Could not extract metadata: {e}")
        
        return metadata
    
    def _parse_core_properties(self, xml_content: str) -> Dict[str, str]:
        properties = {}
        try:
            root = ET.fromstring(xml_content)
            cp_ns = 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties'
            dc_ns = 'http://purl.org/dc/elements/1.1/'
            dcterms_ns = 'http://purl.org/dc/terms/'
            prop_mapping = {
                f'{{{dc_ns}}}title': 'title',
                f'{{{dc_ns}}}creator': 'author',
                f'{{{dc_ns}}}subject': 'subject',
                f'{{{dc_ns}}}description': 'description',
                f'{{{dc_ns}}}publisher': 'publisher',
                f'{{{dc_ns}}}contributor': 'contributor',
                f'{{{dc_ns}}}date': 'date',
                f'{{{dcterms_ns}}}created': 'created',
                f'{{{dcterms_ns}}}modified': 'modified',
                f'{{{cp_ns}}}keywords': 'keywords',
                f'{{{cp_ns}}}category': 'category',
                f'{{{cp_ns}}}version': 'version',
                f'{{{cp_ns}}}contentStatus': 'content_status',
                f'{{{cp_ns}}}language': 'language'
            }
            
            for elem in root:
                if elem.tag in prop_mapping:
                    properties[prop_mapping[elem.tag]] = elem.text
        
        except Exception as e:
            print(f"Warning: Could not parse core properties: {e}")
        
        return properties
    
    def _parse_app_properties(self, xml_content: str) -> Dict[str, str]:
        properties = {}
        try:
            root = ET.fromstring(xml_content)
            app_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'
            
            prop_mapping = {
                f'{{{app_ns}}}Application': 'application',
                f'{{{app_ns}}}DocSecurity': 'security',
                f'{{{app_ns}}}ScaleCrop': 'scale_crop',
                f'{{{app_ns}}}HeadingPairs': 'heading_pairs',
                f'{{{app_ns}}}TitlesOfParts': 'titles_of_parts',
                f'{{{app_ns}}}Company': 'company',
                f'{{{app_ns}}}LinksUpToDate': 'links_up_to_date',
                f'{{{app_ns}}}HyperlinksChanged': 'hyperlinks_changed',
                f'{{{app_ns}}}AppVersion': 'app_version',
                f'{{{app_ns}}}Pages': 'pages',
                f'{{{app_ns}}}Words': 'words',
                f'{{{app_ns}}}Characters': 'characters',
                f'{{{app_ns}}}CharactersWithSpaces': 'characters_with_spaces',
                f'{{{app_ns}}}Lines': 'lines',
                f'{{{app_ns}}}Paragraphs': 'paragraphs',
                f'{{{app_ns}}}TotalTime': 'total_time',
                f'{{{app_ns}}}SharedDoc': 'shared_doc'
            }
            
            for elem in root:
                if elem.tag in prop_mapping:
                    properties[prop_mapping[elem.tag]] = elem.text
        
        except Exception as e:
            print(f"Warning: Could not parse app properties: {e}")
        
        return properties
    
    def _parse_custom_properties(self, xml_content: str) -> Dict[str, str]:
        properties = {}
        
        try:
            root = ET.fromstring(xml_content)
            cp_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/custom-properties'
            vt_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes'
            
            for prop in root.findall(f'.//{{{cp_ns}}}property'):
                name = prop.get('name')
                if name:
                    value_elem = prop.find(f'.//{{{vt_ns}}}lpwstr')
                    if value_elem is not None:
                        properties[name] = value_elem.text
                    else:
                        value_elem = prop.find(f'.//{{{vt_ns}}}i4')
                        if value_elem is not None:
                            properties[name] = value_elem.text
                        else:
                            value_elem = prop.find(f'.//{{{vt_ns}}}r8')
                            if value_elem is not None:
                                properties[name] = value_elem.text
                            else:
                                value_elem = prop.find(f'.//{{{vt_ns}}}bool')
                                if value_elem is not None:
                                    properties[name] = value_elem.text
                                else:
                                    value_elem = prop.find(f'.//{{{vt_ns}}}filetime')
                                    if value_elem is not None:
                                        properties[name] = value_elem.text
        
        except Exception as e:
            print(f"Warning: Could not parse custom properties: {e}")
        
        return properties
    
    def analyze_document_structure(self, filepath: str) -> Dict[str, Any]:    
        structure = {
            'sections': [],
            'paragraphs': [],
            'tables': [],
            'images': [],
            'headers': [],
            'footers': [],
            'footnotes': [],
            'endnotes': [],
            'styles': [],
            'hyperlinks': []
        }
        try:
            from docx import Document
            doc = Document(filepath)
            for i, section in enumerate(doc.sections):
                sect_info = {
                    'index': i,
                    'start_page': i + 1,
                    'page_width': section.page_width.inches,
                    'page_height': section.page_height.inches,
                    'orientation': 'landscape' if section.orientation.value == 1 else 'portrait',
                    'margins': {
                        'left': section.left_margin.inches,
                        'right': section.right_margin.inches,
                        'top': section.top_margin.inches,
                        'bottom': section.bottom_margin.inches
                    }
                }
                structure['sections'].append(sect_info)
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    para_info = {
                        'index': i,
                        'text_length': len(para.text),
                        'style': para.style.name if para.style else 'Normal',
                        'alignment': str(para.alignment) if para.alignment else 'LEFT',
                        'has_runs': len(para.runs) > 0,
                        'first_10_words': ' '.join(para.text.split()[:10])
                    }
                    structure['paragraphs'].append(para_info)
            for i, table in enumerate(doc.tables):
                table_info = {
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'total_cells': len(table.rows) * len(table.columns)
                }
                structure['tables'].append(table_info)
            hyperlinks = self._extract_hyperlinks(filepath)
            structure['hyperlinks'] = hyperlinks
            try:
                with zipfile.ZipFile(filepath) as docx:
                    image_count = sum(1 for name in docx.namelist() 
                                    if name.startswith('word/media/') 
                                    and name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')))
                    structure['images'] = [{'count': image_count}]
            except:
                pass
            metadata = self.extract_metadata(filepath)
            if 'document_properties' in metadata:
                stats = metadata['document_properties']
                structure['statistics'] = {
                    'words': stats.get('words', 0),
                    'characters': stats.get('characters', 0),
                    'pages': stats.get('pages', 0),
                    'paragraphs': stats.get('paragraphs', 0)
                }
        
        except Exception as e:
            print(f"Warning: Could not analyze document structure: {e}")
        
        return structure
    
    def _extract_hyperlinks(self, filepath: str) -> List[Dict[str, str]]:
        hyperlinks = []
        try:
            with zipfile.ZipFile(filepath) as docx:
                if 'word/document.xml' in docx.namelist():
                    xml_content = docx.read('word/document.xml').decode('utf-8')
                    hyperlink_pattern = r'<w:hyperlink[^>]*r:id="([^"]*)"[^>]*>'
                    matches = re.findall(hyperlinks)
                    for rel_id in matches:
                        if 'word/_rels/document.xml.rels' in docx.namelist():
                            rels_content = docx.read('word/_rels/document.xml.rels').decode('utf-8')
                            rel_pattern = f'<Relationship Id="{rel_id}"[^>]*Target="([^"]*)"'
                            rel_match = re.search(rel_pattern, rels_content)
                            
                            if rel_match:
                                target = rel_match.group(1)
                                hyperlinks.append({
                                    'id': rel_id,
                                    'target': target,
                                    'type': 'external' if target.startswith('http') else 'internal'
                                })
        
        except Exception as e:
            print(f"Warning: Could not extract hyperlinks: {e}")
    
    def validate_docx(self, filepath: str) -> Dict[str, Any]:
        validation = {
            'is_valid': False,
            'errors': [],
            'warnings': [],
            'file_info': {}
        }
        try:
            if not Path(filepath).exists():
                validation['errors'].append('File does not exist')
                return validation


def extract_docx_as_zip(filepath: str, extract_to: str = None) -> str:
    if extract_to is None:
        extract_to = tempfile.mkdtemp(prefix='docx_extract_')
        
    try:
        with zipfile.ZipFile(filepath, 'r') as docx:
            docx.extractall(extract_to)
            
        return extract_to
    except Exception as e:
        raise Exception(f"Failed to extract DOCX: {e}")
    
        
    
